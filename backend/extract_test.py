import docx
import os

DOC_PATH = r"C:\Rommaana_C\WixResumes\Candidate_Resumes\_Job Descriptions - Rommaana Sep 25 (1).docx"

if not os.path.exists(DOC_PATH):
    print(f"File not found: {DOC_PATH}")
else:
    doc = docx.Document(DOC_PATH)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text[:100]}...") # Print first 100 chars
